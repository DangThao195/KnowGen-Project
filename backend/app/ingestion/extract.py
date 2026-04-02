import os
import json
import requests
from typing import List, Dict, Any
from pydantic import BaseModel, Field
from docx import Document as DocxDocument
import pymupdf4llm  # Extract PDF to Markdown

from dotenv import load_dotenv
load_dotenv(os.path.join(os.path.dirname(os.getcwd()), '.env'))


# 1. CANONICAL SCHEMA 
class Document(BaseModel):
    id: str
    source_type: str = Field(description="pdf, docx, or notion")
    title: str
    content: str  
    metadata: Dict[str, Any] = Field(default_factory=dict)

# 2. ADAPTERS
def extract_pdf_to_md(file_path: str) -> Document:
    """
    Extract PDF to Markdown.
    """
    print(f"Extracting PDF (Markdown): {file_path}")
    
    md_text = pymupdf4llm.to_markdown(file_path)
    filename = os.path.basename(file_path)
    
    return Document(
        id=filename,
        source_type="pdf",
        title=filename,
        content=md_text.strip(),
        metadata={"path": file_path}
    )

def extract_docx_to_md(file_path: str) -> Document:
    """
    Extract DOCX to Markdown.
    """
    print(f"Extracting DOCX (Markdown): {file_path}")
    doc = DocxDocument(file_path)
    
    md_lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        style_name = para.style.name.lower()
        
        # Mapping Word Styles to Markdown
        if 'heading 1' in style_name:
            md_lines.append(f"# {text}")
        elif 'heading 2' in style_name:
            md_lines.append(f"## {text}")
        elif 'heading 3' in style_name:
            md_lines.append(f"### {text}")
        elif 'heading' in style_name:
            md_lines.append(f"#### {text}")
        elif 'list bullet' in style_name:
            md_lines.append(f"- {text}")
        elif 'list number' in style_name:
            md_lines.append(f"1. {text}")
        else:
            md_lines.append(text)  # Plain Text 
            
    md_text = "\n\n".join(md_lines)
    filename = os.path.basename(file_path)
    
    return Document(
        id=filename,
        source_type="docx",
        title=filename,
        content=md_text.strip(),
        metadata={"path": file_path}
    )

def extract_notion_to_md(page_id: str, token: str) -> Document:
    """
    Extract Notion to Markdown.
    """
    print(f"Extracting Notion (Markdown): {page_id}")
    headers = {
        "Authorization": f"Bearer {token}",
        "Notion-Version": "2022-06-28"
    }
    
    # Get Title Page
    page_url = f"https://api.notion.com/v1/pages/{page_id}"
    page_res = requests.get(page_url, headers=headers)
    page_res.raise_for_status()
    page_data = page_res.json()
    
    title = "Untitled"
    for prop in page_data.get("properties", {}).values():
        if prop.get("type") == "title" and prop["title"]:
            title = prop["title"][0]["plain_text"]
            break

    # Get List Blocks (Content)
    blocks_url = f"https://api.notion.com/v1/blocks/{page_id}/children"
    blocks_res = requests.get(blocks_url, headers=headers)
    blocks_res.raise_for_status()
    blocks = blocks_res.json().get("results", [])
    
    md_lines = []
    for block in blocks:
        b_type = block.get("type")
        if not b_type or b_type not in block:
            continue
            
        # Join plain text from array rich_text
        rich_text = block[b_type].get("rich_text", [])
        raw_text = "".join([t.get("plain_text", "") for t in rich_text])
        if not raw_text and b_type != "divider":
            continue
            
        # Mapping Notion Block Type to Markdown
        if b_type == "heading_1":
            md_lines.append(f"# {raw_text}")
        elif b_type == "heading_2":
            md_lines.append(f"## {raw_text}")
        elif b_type == "heading_3":
            md_lines.append(f"### {raw_text}")
        elif b_type == "bulleted_list_item":
            md_lines.append(f"- {raw_text}")
        elif b_type == "numbered_list_item":
            md_lines.append(f"1. {raw_text}")
        elif b_type == "quote":
            md_lines.append(f"> {raw_text}")
        elif b_type == "code":
            lang = block[b_type].get("language", "")
            md_lines.append(f"```{lang}\n{raw_text}\n```")
        elif b_type == "divider":
            md_lines.append("---")
        else:
            md_lines.append(raw_text)
            
    md_text = "\n\n".join(md_lines)
    
    return Document(
        id=page_id,
        source_type="notion",
        title=title,
        content=md_text.strip(),
        metadata={"url": page_data.get("url")}
    )

# 3. PIPELINE CHÍNH ĐỂ TEST
def run_extractor_demo():
    print("Starting ETL Pipeline...")
    documents: List[Document] = []
    
    sample_path = r"H:\\ĐH 1,2\\BT CNXHKH.docx"
    
    # 1. Scan file Local
    files_to_process = [sample_path] if os.path.isfile(sample_path) else [os.path.join(sample_path, f) for f in os.listdir(sample_path)] if os.path.isdir(sample_path) else []
    
    for file_path in files_to_process:
        filename = os.path.basename(file_path)
        try:
            if filename.endswith(".pdf"):
                documents.append(extract_pdf_to_md(file_path))
            elif filename.endswith(".docx"):
                documents.append(extract_docx_to_md(file_path))
        except Exception as e:
            print(f"Error extracting {filename}: {e}")

    # 2. Scan Notion
    notion_token = os.getenv("NOTION_TOKEN")
    notion_page_id = os.getenv("NOTION_TEST_PAGE_ID")
    
    if notion_token and notion_page_id:
        try:
            documents.append(extract_notion_to_md(notion_page_id, notion_token))
        except Exception as e:
            print(f"Error extracting Notion: {e}")
    else:
        print("Skip Notion: Need to set NOTION_TOKEN and NOTION_TEST_PAGE_ID in .env")

    # 3. Save JSON
    if documents:
        out_file = "extracted_metadata.json"
        with open(out_file, "w", encoding="utf-8") as f:
            json.dump([d.model_dump() for d in documents], f, ensure_ascii=False, indent=2)
        print(f"\nDone! Saved {len(documents)} files to '{out_file}'.")
        
        # Preview first 500 characters of the first file to see Markdown
        print("\n--- PREVIEW FIRST FILE (Demo structure) ---")
        preview = documents[0].content[:500] + "...\n(Đã cắt ngắn bớt)"
        print(preview)

if __name__ == "__main__":
    run_extractor_demo()
